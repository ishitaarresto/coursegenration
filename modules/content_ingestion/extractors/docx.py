"""
DOCX extractor using python-docx.

Responsibilities:
- Open the document via docx.Document(path)
- Walk paragraphs and tables in document order (via doc.element.body) to
  reconstruct the reading order that doc.paragraphs / doc.tables would lose
- Preserve heading hierarchy (Heading 1/2/3) as structural markers
- Detect inline and anchored images inside paragraphs by finding w:drawing
  elements; extract bytes via the relationship ID and insert [Image N]
  placeholders at the exact position the image appears in the text flow
- Populate ExtractedContent.full_text and ExtractedContent.images

Note: python-docx does not expose rendered page numbers -- DOCX has no native
page boundary data, so ExtractedContent.pages is always empty for DOCX.
"""

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from modules.content_ingestion.extractors.base import BaseExtractor
from modules.content_ingestion.models import (
    Asset, AssetType, ExtractedContent, ExtractedImage,
)


def _iter_body_blocks(doc):
    """Yield paragraphs and tables from the document body in document order.

    doc.paragraphs and doc.tables are separate flat lists and lose interleaving.
    Iterating doc.element.body directly preserves the original reading order.
    """
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, doc)


def _extract_para_images(
    para: DocxParagraph,
    doc: Document,
    start_index: int,
) -> list[ExtractedImage]:
    """Extract all images from a paragraph's w:drawing elements.

    Images in DOCX are embedded as relationship blobs. Both inline drawings
    (wp:inline) and anchored drawings (wp:anchor) are handled.
    Dimensions are read from the wp:extent element (EMU units) and converted
    to pixels at 96 DPI: pixels = EMU / 9525.
    """
    images: list[ExtractedImage] = []
    idx = start_index

    for drawing in para._element.findall(".//" + qn("w:drawing")):
        blip = drawing.find(".//" + qn("a:blip"))
        if blip is None:
            continue
        r_embed = blip.get(qn("r:embed"))
        if not r_embed:
            continue
        try:
            image_part = doc.part.related_parts[r_embed]
            blob: bytes = image_part.blob
            mime: str = image_part.content_type or ""

            w_px = h_px = None
            extent = drawing.find(".//" + qn("wp:extent"))
            if extent is not None:
                cx = int(extent.get("cx", 0))  # EMU
                cy = int(extent.get("cy", 0))
                if cx and cy:
                    w_px = round(cx / 9525)   # 914400 EMU/inch ÷ 96 DPI = 9525
                    h_px = round(cy / 9525)

            images.append(ExtractedImage(
                index=idx,
                image_bytes=blob,
                width=w_px,
                height=h_px,
                mime_type=mime,
            ))
            idx += 1
        except (KeyError, AttributeError, Exception):
            continue

    return images


class DocxExtractor(BaseExtractor):
    """Extracts text, structure, and embedded images from DOCX files."""

    def __init__(self, extract_images: bool = True) -> None:
        self.extract_images = extract_images

    def can_handle(self, asset: Asset) -> bool:
        return asset.asset_type == AssetType.DOCX

    def extract(self, asset: Asset) -> ExtractedContent:
        """Walk paragraphs and tables; insert [Image N] placeholders at image positions."""
        self._validate_file(asset)
        result = ExtractedContent(asset=asset)

        try:
            doc = Document(asset.file_path)
        except Exception as exc:
            result.extraction_errors.append(f"Document open failed: {exc}")
            return result

        parts: list[str] = []
        all_images: list[ExtractedImage] = []
        img_index = 0

        for block in _iter_body_blocks(doc):
            if isinstance(block, DocxParagraph):
                # -- Images inside this paragraph ---------------------------
                if self.extract_images:
                    para_images = _extract_para_images(block, doc, img_index)
                    for img in para_images:
                        parts.append(f"[Image {img.index}]")
                    all_images.extend(para_images)
                    img_index += len(para_images)

                # -- Paragraph text -----------------------------------------
                text = block.text.strip()
                if not text:
                    continue
                style_name = block.style.name if block.style else ""
                if "Heading 1" in style_name:
                    parts.append(f"# {text}")
                elif "Heading 2" in style_name:
                    parts.append(f"## {text}")
                elif "Heading 3" in style_name:
                    parts.append(f"### {text}")
                else:
                    parts.append(text)

            elif isinstance(block, DocxTable):
                for row in block.rows:
                    # python-docx repeats the same cell object for merged cells;
                    # deduplicate by skipping consecutive identical values.
                    seen: list[str] = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if not seen or cell_text != seen[-1]:
                            seen.append(cell_text)
                    row_text = " | ".join(c for c in seen if c)
                    if row_text:
                        parts.append(row_text)

        result.full_text = "\n".join(parts)
        result.images = all_images

        props = doc.core_properties
        result.title = props.title or ""
        result.author = props.author or ""
        result.doc_metadata = {
            k: v for k, v in {
                "subject": props.subject,
                "keywords": props.keywords,
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
            }.items() if v
        }

        return result
